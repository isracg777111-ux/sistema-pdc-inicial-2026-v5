import docx
import shutil
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

try:
    # First make a copy to avoid read-lock issues if any
    shutil.copyfile('inicial.docx', 'inicial_temp_2.docx')
    doc = docx.Document('inicial_temp_2.docx')
    
    # Text to match to apply formatting
    marker_tags = ['campos_areas', 'perfil_de_salida', 'momentos_formativos', 'recursos_semana', 'criterios_de_evaluacion', 'condicion', 'adaptaciones', 'distrito_educativo', 'director', 'maestro', 'trimestre', 'duracion', 'unidad_educativa', 'año_de_escolaridad', 'numero_pdc', 'nivel']
    
    def applies_to(text):
        if not text: return False
        return any(tag in text for tag in marker_tags)
        
    for para in doc.paragraphs:
        if applies_to(para.text):
            para.alignment = WD_ALIGN_PARAGRAPH.LEFT
            para.paragraph_format.line_spacing = 1.15
            for run in para.runs:
                run.font.name = 'Arial'
                run.font.size = Pt(9)
                
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    if applies_to(para.text):
                        para.alignment = WD_ALIGN_PARAGRAPH.LEFT
                        para.paragraph_format.line_spacing = 1.15
                        for run in para.runs:
                            run.font.name = 'Arial'
                            run.font.size = Pt(9)
                            
    doc.save('inicial_formateado.docx')
    print("SUCCESS: inicial_formateado.docx created")
except Exception as e:
    print(f"Error: {e}")
