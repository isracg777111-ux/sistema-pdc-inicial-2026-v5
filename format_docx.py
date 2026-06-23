import docx
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

def format_docx(filepath):
    doc = docx.Document(filepath)
    
    # Apply to paragraphs
    for para in doc.paragraphs:
        if "{{" in para.text:
            para.alignment = WD_ALIGN_PARAGRAPH.LEFT
            para.paragraph_format.line_spacing = 1.15
            for run in para.runs:
                run.font.name = 'Arial'
                run.font.size = Pt(9)
                
    # Apply to tables
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    if "{{" in para.text:
                        para.alignment = WD_ALIGN_PARAGRAPH.LEFT
                        para.paragraph_format.line_spacing = 1.15
                        for run in para.runs:
                            run.font.name = 'Arial'
                            run.font.size = Pt(9)
                            
    doc.save('inicial.docx')
    print("Formatting applied successfully!")

try:
    format_docx('inicial.docx')
except Exception as e:
    print(f"Error: {e}")
